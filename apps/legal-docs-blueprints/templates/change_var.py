#!/usr/bin/env python3
"""
Script para RENOMBRAR variables PRESERVANDO el formato del documento
(negritas, cursivas, tamaños de fuente, etc.)
"""

from docx import Document
import re

# Documento de entrada
ARCHIVO_ENTRADA = 'templates/contrato_uso_carro_usado.docx'

# Documento de salida
ARCHIVO_SALIDA = 'templates/contrato_uso_carro_usado_renombrado.docx'

# ==================== MAPEO DE VARIABLES ====================
RENOMBRAR = {
    '{date_day}': '{diaTexto}',
    '{date_month}': '{mesTexto}',
    '{date_year}': '{anoTexto}',
    '{contract_day}': '{diaTexto}',
    '{contract_month}': '{mesTexto}',
    '{contract_year}': '{anoTexto}',
    '{contract_start_date}': '{fechaInicioContrato}',
    '{contract_end_day}': '{diaTextoVencimiento}',
    '{contract_end_month}': '{mesTextoVencimiento}',
    '{contract_end_year}': '{anoTextoVencimiento}',
    '{debtor_name}': '{nombreCompleto}',
    '{client_name}': '{nombreCompleto}',
    
    '{debtor_dpi_number}': '{dpi}',
    '{client_cui}': '{dpiTexto}',  # CUI = DPI en Guatemala
    
    '{debtor_dpi_letters}': '{dpiTexto}',
    '{client_age}': '{edad}',
    '{client_marital_status_gendered}': '{estadoCivil}',
    '{client_occupation}': '{profesion}',
    '{client_nationality_gendered}': '{nacionalidad}',
    '{client_address}': '{direccion}',
    '{date_day}': '{dia}',
    '{date_month}': '{mes}',
    '{date_year}': '{ano}',
    '{debtor_name}': '{nombreCompleto}',
    '{debtor_dpi_number}': '{dpi}',
    '{debtor_dpi_letters}': '{dpiTexto}',
    '{vehicle_type}': '{tipoVehiculo}',
    '{vehicle_brand}': '{marcaVehiculo}',
    '{vehicle_color}': '{colorVehiculo}',
    '{vehicle_use}': '{usoVehiculo}',
    '{vehicle_chassis}': '{chasisVehiculo}',
    '{vehicle_fuel}': '{combustibleVehiculo}',
    '{vehicle_engine}': '{motorVehiculo}',
    '{vehicle_series}': '{serieVehiculo}',
    '{vehicle_line}': '{lineaVehiculo}',
    '{vehicle_model}': '{modeloVehiculo}',
    '{vehicle_cc}': '{cm3Vehiculo}',
    '{vehicle_seats}': '{asientosVehiculo}',
    '{vehicle_cylinders}': '{cilindrosVehiculo}',
    '{vehicle_iscv}': '{iscvVehiculo}',
    '{contract_duration_months}': '{plazoTexto}',
     '{user_name}': '{nombreCompleto}',
    '{user_name_clause_a}': '{nombreCompleto}',
    '{user_name_clause_a2}': '{nombreCompleto}',
    '{user_name_clause_b}': '{nombreCompleto}',
    '{user_name_clause_d}': '{nombreCompleto}',
    '{user_name_final}': '{nombreCompleto}',
}


def renombrar_en_runs(runs):
    """
    Renombra variables en los runs preservando el formato.
    Los runs son fragmentos de texto con el mismo formato.
    """
    contador = 0
    
    for run in runs:
        texto_original = run.text
        texto_nuevo = texto_original
        
        # Reemplazar cada variable
        for viejo, nuevo in RENOMBRAR.items():
            if viejo in texto_nuevo:
                texto_nuevo = texto_nuevo.replace(viejo, nuevo)
                contador += 1
        
        # Solo actualizar si hubo cambios
        if texto_original != texto_nuevo:
            run.text = texto_nuevo
    
    return contador


def renombrar_en_parrafos(doc):
    """Renombra las variables en los párrafos preservando formato"""
    contador = 0
    
    for para in doc.paragraphs:
        contador += renombrar_en_runs(para.runs)
    
    return contador


def renombrar_en_tablas(doc):
    """Renombra las variables en las tablas preservando formato"""
    contador = 0
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Procesar párrafos dentro de la celda
                for para in cell.paragraphs:
                    contador += renombrar_en_runs(para.runs)
    
    return contador


def mostrar_cambios():
    """Muestra los cambios que se van a realizar"""
    print("\n📋 Cambios a realizar:\n")
    for viejo, nuevo in RENOMBRAR.items():
        print(f"  {viejo:30} → {nuevo}")
    print()


def main():
    print("\n" + "="*60)
    print("RENOMBRAR VARIABLES PRESERVANDO FORMATO")
    print("="*60)
    
    # Mostrar cambios
    mostrar_cambios()
    
    # Cargar documento
    print(f"📄 Cargando documento: {ARCHIVO_ENTRADA}")
    try:
        doc = Document(ARCHIVO_ENTRADA)
    except FileNotFoundError:
        print(f"❌ Error: No se encontró '{ARCHIVO_ENTRADA}'")
        print("   Verifica que el archivo existe en esta carpeta")
        return
    
    print("🔄 Renombrando variables (preservando formato)...\n")
    
    # Renombrar en párrafos
    total_parrafos = renombrar_en_parrafos(doc)
    print(f"  ✓ Párrafos: {total_parrafos} renombramientos")
    
    # Renombrar en tablas
    total_tablas = renombrar_en_tablas(doc)
    print(f"  ✓ Tablas: {total_tablas} renombramientos")
    
    total = total_parrafos + total_tablas
    
    # Guardar documento
    print(f"\n💾 Guardando documento: {ARCHIVO_SALIDA}")
    doc.save(ARCHIVO_SALIDA)
    
    print(f"\n✅ ¡Completado! Total: {total} renombramientos")
    print(f"📁 Archivo guardado: {ARCHIVO_SALIDA}")
    print(f"   El formato original se ha preservado (negritas, etc.)\n")


if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        print(f"\n❌ Error inesperado: {e}")
        import traceback
        traceback.print_exc()